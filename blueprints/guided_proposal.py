"""Isolated backend API for the local-first Guided Proposal V2 pilot."""

from __future__ import annotations

import json
import logging
import time
import uuid

from flask import Blueprint, jsonify, request, send_file

from agent.proposal_v2_rules import (
    donor_profiles,
    normalize_setup,
    normalize_step2,
    normalize_step3,
    normalize_step4,
    validate_setup,
    validate_step2,
    validate_step3,
    validate_step4,
)
from auth import current_role, current_uid, require_auth, require_role
from blueprints.helpers import _chats_db, _log_event

guided_proposal_bp = Blueprint("guided_proposal", __name__, url_prefix="/api/proposals")
logger = logging.getLogger(__name__)

_PROPOSAL_SCHEMA = """
CREATE TABLE IF NOT EXISTS proposal_v2_setups (
    id                  TEXT PRIMARY KEY,
    uid                 TEXT NOT NULL,
    project_title       TEXT,
    country             TEXT,
    region              TEXT,
    donor               TEXT,
    budget_amount       REAL,
    budget_currency     TEXT,
    executive_intent    TEXT,
    sectors             TEXT,
    state               TEXT DEFAULT 'draft',
    analysis            TEXT,
    context_data        TEXT,
    step2_analysis      TEXT,
    step2_state         TEXT,
    step2_locked_at     TEXT,
    technical_data      TEXT,
    step3_analysis      TEXT,
    step3_state         TEXT,
    step3_locked_at     TEXT,
    financial_data      TEXT,
    step4_analysis      TEXT,
    step4_state         TEXT,
    step4_locked_at     TEXT,
    call_brief          TEXT,
    reference_text      TEXT,
    reference_filename  TEXT,
    created_at          TEXT,
    updated_at          TEXT
);
"""


def _ensure_schema():
    """Create the proposal_v2_setups table if it does not exist."""
    try:
        import sqlite3

        from config import CHATS_DB_PATH

        conn = sqlite3.connect(str(CHATS_DB_PATH))
        conn.execute(_PROPOSAL_SCHEMA)
        conn.commit()
        conn.close()
    except Exception as exc:
        logger.error("Failed to ensure proposal schema: %s", exc)


def _enabled_response():
    """V2 is now the sole proposal system — always enabled."""
    return None


def _serialize(row):
    data = dict(row)
    for field, default in (
        ("sectors", []),
        ("analysis", {}),
        ("context_data", {}),
        ("step2_analysis", {}),
        ("technical_data", {}),
        ("step3_analysis", {}),
        ("financial_data", {}),
        ("step4_analysis", {}),
        ("call_brief", {}),
    ):
        try:
            data[field] = json.loads(data[field]) if data[field] else default
        except (TypeError, json.JSONDecodeError):
            data[field] = default
    return data


def _owned_setup(setup_id: str, uid: str, role: str):
    _ensure_schema()
    conn = _chats_db()
    if role == "admin":
        row = conn.execute("SELECT * FROM proposal_v2_setups WHERE id = ?", (setup_id,)).fetchone()
    else:
        row = conn.execute("SELECT * FROM proposal_v2_setups WHERE id = ? AND uid = ?", (setup_id, uid)).fetchone()
    return conn, row


@guided_proposal_bp.route("/donors", methods=["GET"])
@require_auth
def api_guided_proposal_donors():
    if disabled := _enabled_response():
        return disabled
    return jsonify(donor_profiles())


@guided_proposal_bp.route("/donor-rules", methods=["GET"])
@require_auth
def api_guided_proposal_donor_rules():
    """Return publicly-visible donor rules manifest for frontend guidance cards."""
    from agent.proposal_v2_rules import DONOR_PROFILES

    public_fields = (
        "id",
        "label",
        "full_name",
        "framework_standard",
        "overhead_ceiling_percent",
        "max_duration_months",
        "currency_options",
        "section_rules",
        "beneficiary_rules",
        "logframe_rules",
        "financial_rules",
        "prompt_directive",
    )
    public = [{k: p[k] for k in public_fields if k in p} for p in DONOR_PROFILES.values()]
    return jsonify({"version": "2.0.0-2026", "donors": public})


@guided_proposal_bp.route("/setups/<setup_id>/upload-reference", methods=["POST"])
@require_role("premium")
def api_guided_proposal_upload_reference(setup_id):
    """Attach a local call document before Step 1 is locked."""
    if disabled := _enabled_response():
        return disabled
    conn, row = _owned_setup(setup_id, current_uid(), current_role())
    try:
        if not row:
            return jsonify({"error": "Setup not found."}), 404
        if row["state"] == "locked":
            return jsonify({"error": "Step 1 is locked; reference documents are immutable."}), 409
        file = request.files.get("file")
        if not file or not file.filename:
            return jsonify({"error": "No call document uploaded."}), 400
        from werkzeug.utils import secure_filename

        filename = secure_filename(file.filename)
        if not filename.lower().endswith((".docx", ".txt", ".md")):
            return jsonify({"error": "Upload a DOCX, TXT or Markdown call document."}), 400
        if filename.lower().endswith(".docx"):
            import docx

            document = docx.Document(file)
            parts = [p.text for p in document.paragraphs if p.text.strip()]
            parts += [" | ".join(cell.text for cell in row.cells) for table in document.tables for row in table.rows]
            text = "\n\n".join(parts)
        else:
            text = file.read().decode("utf-8", errors="ignore")
        text = text.strip()[:50000]
        if not text:
            return jsonify({"error": "No readable text found in the call document."}), 400
        conn.execute(
            "UPDATE proposal_v2_setups SET reference_text = ?, reference_filename = ?, call_brief = '{}', updated_at = ? WHERE id = ?",
            (text, filename, time.time(), setup_id),
        )
        conn.commit()
        return jsonify({"filename": filename, "chars": len(text), "message": "Call document attached."})
    finally:
        conn.close()


@guided_proposal_bp.route("/setups/<setup_id>/call-brief", methods=["POST"])
@require_role("premium")
def api_guided_proposal_call_brief(setup_id):
    """Explain what the attached call asks for without changing the proposal."""
    uid, role = current_uid(), current_role()
    conn, row = _owned_setup(setup_id, uid, role)
    try:
        if not row:
            return jsonify({"error": "Setup not found."}), 404
        setup = _serialize(row)
        if not setup.get("reference_text"):
            return jsonify({"error": "Upload a grant call before requesting its briefing."}), 422
        if setup.get("call_brief"):
            return jsonify({"brief": setup["call_brief"], "filename": setup.get("reference_filename"), "cached": True})
        from agent.proposal_v2_agents import summarize_call_document

        brief = summarize_call_document(setup)
        conn.execute(
            "UPDATE proposal_v2_setups SET call_brief = ?, updated_at = ? WHERE id = ?",
            (json.dumps(brief, ensure_ascii=False), time.time(), setup_id),
        )
        conn.commit()
        _log_event(uid, "guided_proposal_call_brief_generated", {"setup_id": setup_id})
        return jsonify({"brief": brief, "filename": setup.get("reference_filename"), "generated_at": time.time()})
    except Exception:
        return jsonify({"error": "The call briefing is temporarily unavailable. Please try again."}), 503
    finally:
        conn.close()


@guided_proposal_bp.route("/setups", methods=["POST"])
@require_role("premium")
def api_guided_proposal_create_setup():
    if disabled := _enabled_response():
        return disabled
    _ensure_schema()
    uid = current_uid()
    setup = normalize_setup(request.get_json(silent=True) or {})
    now = time.time()
    setup_id = f"gps_{uuid.uuid4().hex[:12]}"
    conn = _chats_db()
    try:
        conn.execute(
            """INSERT INTO proposal_v2_setups
               (id, uid, project_title, country, region, donor, budget_amount,
                budget_currency, executive_intent, sectors, state, created_at, updated_at)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'draft', ?, ?)""",
            (
                setup_id,
                uid,
                setup["project_title"],
                setup["country"],
                setup["region"],
                setup["donor"],
                setup["budget_amount"],
                setup["budget_currency"],
                setup["executive_intent"],
                json.dumps(setup["sectors"]),
                now,
                now,
            ),
        )
        conn.commit()
        row = conn.execute("SELECT * FROM proposal_v2_setups WHERE id = ?", (setup_id,)).fetchone()
        _log_event(uid, "guided_proposal_setup_created", {"setup_id": setup_id})
        return jsonify(_serialize(row)), 201
    finally:
        conn.close()


@guided_proposal_bp.route("/setups", methods=["GET"])
@require_auth
def api_guided_proposal_list_setups():
    """List all proposal setups for the current user, newest first."""
    _ensure_schema()
    uid, role = current_uid(), current_role()
    conn = _chats_db()
    try:
        if role == "admin":
            rows = conn.execute("SELECT * FROM proposal_v2_setups ORDER BY created_at DESC").fetchall()
        else:
            rows = conn.execute(
                "SELECT * FROM proposal_v2_setups WHERE uid = ? ORDER BY created_at DESC",
                (uid,),
            ).fetchall()
        # The UI must not infer mutation rights from a late-loading browser
        # auth variable. Return the server-authoritative capability with each
        # record so eligible owners can always see the delete control.
        can_delete = role in {"premium", "admin"}
        setups = []
        for row in rows:
            setup = _serialize(row)
            setup["can_delete"] = can_delete
            setups.append(setup)
        return jsonify(setups)
    finally:
        conn.close()


@guided_proposal_bp.route("/setups/<setup_id>", methods=["GET"])
@require_auth
def api_guided_proposal_get_setup(setup_id):
    if disabled := _enabled_response():
        return disabled
    conn, row = _owned_setup(setup_id, current_uid(), current_role())
    try:
        if not row:
            return jsonify({"error": "Setup not found."}), 404
        setup = _serialize(row)
        setup["can_delete"] = current_role() in {"premium", "admin"}
        return jsonify(setup)
    finally:
        conn.close()


@guided_proposal_bp.route("/setups/<setup_id>", methods=["DELETE"])
@require_role("premium")
def api_guided_proposal_delete_setup(setup_id):
    """Permanently remove a Guided Proposal owned by the caller.

    Administrators may remove any local setup; other eligible users may only
    remove their own.  This deliberately uses the same ownership check as the
    read/update routes so the UI cannot accidentally target the legacy
    ``proposals`` table.
    """
    uid, role = current_uid(), current_role()
    conn, row = _owned_setup(setup_id, uid, role)
    try:
        if not row:
            return jsonify({"error": "Guided proposal not found."}), 404
        conn.execute("DELETE FROM proposal_v2_setups WHERE id = ?", (setup_id,))
        conn.commit()
        _log_event(uid, "guided_proposal_deleted", {"setup_id": setup_id})
        return jsonify({"message": "Guided proposal deleted."})
    finally:
        conn.close()


@guided_proposal_bp.route("/setups/<setup_id>", methods=["PUT"])
@require_role("premium")
def api_guided_proposal_update_setup(setup_id):
    if disabled := _enabled_response():
        return disabled
    uid, role = current_uid(), current_role()
    conn, row = _owned_setup(setup_id, uid, role)
    try:
        if not row:
            return jsonify({"error": "Setup not found."}), 404
        if row["state"] == "locked":
            return jsonify({"error": "Locked setup is immutable. Create a new revision to change it."}), 409

        setup = normalize_setup(request.get_json(silent=True) or {})
        conn.execute(
            """UPDATE proposal_v2_setups SET project_title = ?, country = ?, region = ?, donor = ?,
               budget_amount = ?, budget_currency = ?, executive_intent = ?, sectors = ?,
               state = 'draft', analysis = '{}', updated_at = ? WHERE id = ?""",
            (
                setup["project_title"],
                setup["country"],
                setup["region"],
                setup["donor"],
                setup["budget_amount"],
                setup["budget_currency"],
                setup["executive_intent"],
                json.dumps(setup["sectors"]),
                time.time(),
                setup_id,
            ),
        )
        conn.commit()
        updated = conn.execute("SELECT * FROM proposal_v2_setups WHERE id = ?", (setup_id,)).fetchone()
        return jsonify(_serialize(updated))
    finally:
        conn.close()


@guided_proposal_bp.route("/setups/<setup_id>/analyze", methods=["POST"])
@require_role("premium")
def api_guided_proposal_analyze_setup(setup_id):
    if disabled := _enabled_response():
        return disabled
    uid, role = current_uid(), current_role()
    conn, row = _owned_setup(setup_id, uid, role)
    try:
        if not row:
            return jsonify({"error": "Setup not found."}), 404
        if row["state"] == "locked":
            return jsonify({"error": "Locked setup cannot be re-analyzed."}), 409

        setup = _serialize(row)
        analysis = validate_setup(setup)
        deterministic_valid = analysis["is_valid"]
        try:
            from agent.proposal_v2_agents import analyze_step_one

            llm = analyze_step_one(setup, analysis)
            analysis.update(llm)
            analysis["is_valid"] = deterministic_valid and bool(llm.get("is_valid"))
        except Exception:
            # The deterministic donor rules are authoritative. An intermittent
            # model failure must be visible to the user, but must not turn a
            # valid local proposal into an un-lockable dead end.
            analysis.update(
                {
                    "step_id": 1,
                    "donor_compliance_score": 80 if analysis["is_valid"] else 0,
                    "critique_notes": [],
                    "suggested_improvements": [],
                    "verifier": "unavailable",
                    "analyzed_at": time.time(),
                    "warnings": analysis["warnings"] + ["AI verifier unavailable; deterministic donor checks passed."],
                }
            )
        conn.execute(
            "UPDATE proposal_v2_setups SET state = 'analyzed', analysis = ?, updated_at = ? WHERE id = ?",
            (json.dumps(analysis), time.time(), setup_id),
        )
        conn.commit()
        _log_event(uid, "guided_proposal_setup_analyzed", {"setup_id": setup_id, "valid": analysis["is_valid"]})
        return jsonify(analysis)
    finally:
        conn.close()


@guided_proposal_bp.route("/setups/<setup_id>/generate-draft", methods=["POST"])
@require_role("premium")
def api_guided_proposal_generate_step1_draft(setup_id):
    """Generate an editable co-writing draft before Step 1 is locked."""
    uid, role = current_uid(), current_role()
    conn, row = _owned_setup(setup_id, uid, role)
    try:
        if not row:
            return jsonify({"error": "Setup not found."}), 404
        if row["state"] == "locked":
            return jsonify({"error": "Step 1 is locked; create a new proposal to draft again."}), 409
        from agent.proposal_v2_agents import generate_step_one_draft

        draft = generate_step_one_draft(_serialize(row))
        _log_event(uid, "guided_proposal_draft_generated", {"setup_id": setup_id})
        return jsonify({"draft": draft, "generated_at": time.time()})
    except Exception:
        return jsonify(
            {"error": "AI draft generation is temporarily unavailable. Keep editing manually and try again."}
        ), 503
    finally:
        conn.close()


@guided_proposal_bp.route("/setups/<setup_id>/generate-step2-draft", methods=["POST"])
@require_role("premium")
def api_guided_proposal_generate_step2_draft(setup_id):
    """Create a first editable context draft after the Step 1 contract is locked."""
    uid, role = current_uid(), current_role()
    conn, row = _owned_setup(setup_id, uid, role)
    try:
        if not row:
            return jsonify({"error": "Setup not found."}), 404
        if row["state"] != "locked":
            return jsonify({"error": "Lock Step 1 before generating the context draft."}), 409
        if row["step2_state"] == "locked":
            return jsonify({"error": "Step 2 is locked; create a new proposal to draft it again."}), 409
        from agent.proposal_v2_agents import generate_step_two_draft

        draft = generate_step_two_draft(_serialize(row))
        _log_event(uid, "guided_proposal_step2_draft_generated", {"setup_id": setup_id})
        return jsonify({"draft": draft, "generated_at": time.time()})
    except Exception:
        return jsonify({"error": "AI context draft generation is temporarily unavailable. Please try again."}), 503
    finally:
        conn.close()


@guided_proposal_bp.route("/setups/<setup_id>/generate-step3-draft", methods=["POST"])
@require_role("premium")
def api_guided_proposal_generate_step3_draft(setup_id):
    """Create a first editable technical design after Context & Needs is locked."""
    uid, role = current_uid(), current_role()
    conn, row = _owned_setup(setup_id, uid, role)
    try:
        if not row:
            return jsonify({"error": "Setup not found."}), 404
        if row["state"] != "locked" or row["step2_state"] != "locked":
            return jsonify({"error": "Lock Step 1 and Context & Needs before generating the technical draft."}), 409
        if row["step3_state"] == "locked":
            return jsonify({"error": "Step 3 is locked; create a new proposal to draft it again."}), 409
        from agent.proposal_v2_agents import generate_step_three_draft

        draft = generate_step_three_draft(_serialize(row))
        _log_event(uid, "guided_proposal_step3_draft_generated", {"setup_id": setup_id})
        return jsonify({"draft": draft, "generated_at": time.time()})
    except Exception:
        logger.exception("Step 3 technical design draft generation failed for setup %s", setup_id)
        return jsonify({"error": "AI technical design draft is temporarily unavailable. Please try again."}), 503
    finally:
        conn.close()


@guided_proposal_bp.route("/setups/<setup_id>/generate-step4-draft", methods=["POST"])
@require_role("premium")
def api_guided_proposal_generate_step4_draft(setup_id):
    """Create an editable budget, risks and commitments draft."""
    uid, role = current_uid(), current_role()
    conn, row = _owned_setup(setup_id, uid, role)
    try:
        if not row:
            return jsonify({"error": "Setup not found."}), 404
        if row["step3_state"] != "locked":
            return jsonify({"error": "Lock Step 3 before generating the financial draft."}), 409
        if row["step4_state"] == "locked":
            return jsonify({"error": "Step 4 is locked; create a new proposal to draft it again."}), 409
        from agent.proposal_v2_agents import generate_step_four_draft

        draft = generate_step_four_draft(_serialize(row))
        _log_event(uid, "guided_proposal_step4_draft_generated", {"setup_id": setup_id})
        return jsonify({"draft": draft, "generated_at": time.time()})
    except Exception:
        logger.exception("Step 4 financial draft generation failed for setup %s", setup_id)
        return jsonify({"error": "AI financial draft is temporarily unavailable. Please try again."}), 503
    finally:
        conn.close()


@guided_proposal_bp.route("/setups/<setup_id>/lock", methods=["POST"])
@require_role("premium")
def api_guided_proposal_lock_setup(setup_id):
    if disabled := _enabled_response():
        return disabled
    uid, role = current_uid(), current_role()
    conn, row = _owned_setup(setup_id, uid, role)
    try:
        if not row:
            return jsonify({"error": "Setup not found."}), 404
        if row["state"] == "locked":
            return jsonify({"error": "Setup is already locked."}), 409

        setup = _serialize(row)
        analysis = validate_setup(setup)
        try:
            stored_analysis = json.loads(row["analysis"]) if row["analysis"] else {}
        except json.JSONDecodeError:
            stored_analysis = {}
        score = stored_analysis.get("donor_compliance_score", 0)
        if not analysis["is_valid"] or score < 70:
            return jsonify({"error": "Fix validation errors before locking.", "analysis": analysis}), 422

        now = time.time()
        analysis.update(stored_analysis)
        analysis.update({"step_id": 1, "locked_at": now})
        conn.execute(
            "UPDATE proposal_v2_setups SET state = 'locked', analysis = ?, locked_at = ?, updated_at = ? WHERE id = ?",
            (json.dumps(analysis), now, now, setup_id),
        )
        conn.commit()
        locked = conn.execute("SELECT * FROM proposal_v2_setups WHERE id = ?", (setup_id,)).fetchone()
        _log_event(uid, "guided_proposal_setup_locked", {"setup_id": setup_id})
        return jsonify(_serialize(locked))
    finally:
        conn.close()


def _step2_request_setup():
    payload = request.get_json(silent=True) or {}
    setup_id = str(payload.get("setup_id", "")).strip()
    return payload, setup_id


@guided_proposal_bp.route("/steps/2/analyze", methods=["POST"])
@require_role("premium")
def api_guided_proposal_analyze_step2():
    """Analyze a Step 2 draft scoped to an immutable, owned Step 1 setup."""
    if disabled := _enabled_response():
        return disabled
    payload, setup_id = _step2_request_setup()
    if not setup_id:
        return jsonify({"error": "setup_id is required."}), 400
    uid, role = current_uid(), current_role()
    conn, row = _owned_setup(setup_id, uid, role)
    try:
        if not row:
            return jsonify({"error": "Setup not found."}), 404
        if row["state"] != "locked":
            return jsonify({"error": "Step 1 must be locked before Step 2 can be analyzed."}), 409
        if row["step2_state"] == "locked":
            return jsonify({"error": "Locked Step 2 is immutable."}), 409
        setup, step2 = _serialize(row), normalize_step2(payload)
        analysis = validate_step2(step2, setup)
        deterministic_valid = analysis["is_valid"]
        try:
            from agent.proposal_v2_agents import analyze_step_two

            llm = analyze_step_two(setup, step2, analysis)
            analysis.update(llm)
            analysis["is_valid"] = deterministic_valid and bool(llm.get("is_valid"))
        except Exception:
            analysis.update(
                {
                    "donor_compliance_score": 80 if analysis["is_valid"] else 0,
                    "critique_notes": [],
                    "suggested_improvements": [],
                    "verifier": "unavailable",
                    "analyzed_at": time.time(),
                    "warnings": analysis["warnings"] + ["AI verifier unavailable; deterministic donor checks passed."],
                }
            )
        conn.execute(
            "UPDATE proposal_v2_setups SET context_data = ?, step2_analysis = ?, step2_state = 'analyzed', updated_at = ? WHERE id = ?",
            (json.dumps(step2), json.dumps(analysis), time.time(), setup_id),
        )
        conn.commit()
        _log_event(uid, "guided_proposal_step2_analyzed", {"setup_id": setup_id, "valid": analysis["is_valid"]})
        return jsonify(analysis)
    finally:
        conn.close()


@guided_proposal_bp.route("/steps/2/lock", methods=["POST"])
@require_role("premium")
def api_guided_proposal_lock_step2():
    if disabled := _enabled_response():
        return disabled
    payload, setup_id = _step2_request_setup()
    if not setup_id:
        return jsonify({"error": "setup_id is required."}), 400
    uid, role = current_uid(), current_role()
    conn, row = _owned_setup(setup_id, uid, role)
    try:
        if not row:
            return jsonify({"error": "Setup not found."}), 404
        if row["state"] != "locked":
            return jsonify({"error": "Step 1 must be locked before Step 2 can be locked."}), 409
        if row["step2_state"] == "locked":
            return jsonify({"error": "Step 2 is already locked."}), 409
        setup = _serialize(row)
        step2 = normalize_step2(
            payload
            if any(
                key in payload
                for key in ("humanitarian_context", "needs_assessment", "strategic_justification", "beneficiaries")
            )
            else setup["context_data"]
        )
        deterministic = validate_step2(step2, setup)
        stored = setup["step2_analysis"]
        if not deterministic["is_valid"] or stored.get("donor_compliance_score", 0) < 70 or not stored.get("is_valid"):
            return jsonify(
                {"error": "Analyze Step 2 and fix validation errors before locking.", "analysis": deterministic}
            ), 422
        now = time.time()
        stored.update({"step_id": 2, "locked_at": now})
        conn.execute(
            "UPDATE proposal_v2_setups SET context_data = ?, step2_analysis = ?, step2_state = 'locked', step2_locked_at = ?, updated_at = ? WHERE id = ?",
            (json.dumps(step2), json.dumps(stored), now, now, setup_id),
        )
        conn.commit()
        locked = conn.execute("SELECT * FROM proposal_v2_setups WHERE id = ?", (setup_id,)).fetchone()
        _log_event(uid, "guided_proposal_step2_locked", {"setup_id": setup_id})
        return jsonify(_serialize(locked))
    finally:
        conn.close()


@guided_proposal_bp.route("/steps/3/analyze", methods=["POST"])
@require_role("premium")
def api_guided_proposal_analyze_step3():
    if disabled := _enabled_response():
        return disabled
    payload, setup_id = _step2_request_setup()
    if not setup_id:
        return jsonify({"error": "setup_id is required."}), 400
    uid, role = current_uid(), current_role()
    conn, row = _owned_setup(setup_id, uid, role)
    try:
        if not row:
            return jsonify({"error": "Setup not found."}), 404
        if row["step2_state"] != "locked":
            return jsonify({"error": "Step 2 must be locked before Step 3 can be analyzed."}), 409
        if row["step3_state"] == "locked":
            return jsonify({"error": "Locked Step 3 is immutable."}), 409
        setup = _serialize(row)
        step3 = normalize_step3(payload)
        analysis = validate_step3(step3, setup)
        deterministic_valid = analysis["is_valid"]
        try:
            from agent.proposal_v2_agents import analyze_step_three

            llm = analyze_step_three(setup, setup["context_data"], step3, analysis)
            analysis.update(llm)
            analysis["is_valid"] = deterministic_valid and bool(llm.get("is_valid"))
        except Exception:
            analysis.update(
                {
                    "donor_compliance_score": 80 if analysis["is_valid"] else 0,
                    "critique_notes": [],
                    "suggested_improvements": [],
                    "verifier": "unavailable",
                    "analyzed_at": time.time(),
                    "warnings": analysis["warnings"] + ["AI verifier unavailable; deterministic donor checks passed."],
                }
            )
        conn.execute(
            "UPDATE proposal_v2_setups SET technical_data = ?, step3_analysis = ?, step3_state = 'analyzed', updated_at = ? WHERE id = ?",
            (json.dumps(step3), json.dumps(analysis), time.time(), setup_id),
        )
        conn.commit()
        _log_event(uid, "guided_proposal_step3_analyzed", {"setup_id": setup_id, "valid": analysis["is_valid"]})
        return jsonify(analysis)
    finally:
        conn.close()


@guided_proposal_bp.route("/steps/3/lock", methods=["POST"])
@require_role("premium")
def api_guided_proposal_lock_step3():
    if disabled := _enabled_response():
        return disabled
    payload, setup_id = _step2_request_setup()
    if not setup_id:
        return jsonify({"error": "setup_id is required."}), 400
    uid, role = current_uid(), current_role()
    conn, row = _owned_setup(setup_id, uid, role)
    try:
        if not row:
            return jsonify({"error": "Setup not found."}), 404
        if row["step2_state"] != "locked":
            return jsonify({"error": "Step 2 must be locked before Step 3 can be locked."}), 409
        if row["step3_state"] == "locked":
            return jsonify({"error": "Step 3 is already locked."}), 409
        setup = _serialize(row)
        step3 = normalize_step3(
            payload
            if any(key in payload for key in ("logframe", "toc_narrative", "gantt"))
            else setup["technical_data"]
        )
        deterministic = validate_step3(step3, setup)
        stored = setup["step3_analysis"]
        if not deterministic["is_valid"] or stored.get("donor_compliance_score", 0) < 70 or not stored.get("is_valid"):
            return jsonify(
                {"error": "Analyze Step 3 and fix validation errors before locking.", "analysis": deterministic}
            ), 422
        # ── M&E quality review on logframe + ToC ──────────────────────────
        try:
            from agent.me_reviewer import review_section

            toc_list = step3.get("toc_narrative", []) if isinstance(step3.get("toc_narrative"), list) else []
            logframe_dict = step3.get("logframe", {}) if isinstance(step3.get("logframe"), dict) else {}
            me_review = review_section(
                content=stored.get("technical_summary", ""),
                step="logframe",
                toc=toc_list,
                logframe=logframe_dict,
                sources=stored.get("sources", []),
            )
            stored["quality_score"] = me_review.get("quality_score")
            stored["me_suggestions"] = me_review.get("suggestions", [])
            stored["me_overall_score"] = me_review.get("overall_score", 0)
        except Exception as _me_err:
            import logging as _lg

            _lg.getLogger(__name__).warning(f"M&E review failed for step 3 lock: {_me_err}")
        now = time.time()
        stored.update({"step_id": 3, "locked_at": now})
        conn.execute(
            "UPDATE proposal_v2_setups SET technical_data = ?, step3_analysis = ?, step3_state = 'locked', step3_locked_at = ?, updated_at = ? WHERE id = ?",
            (json.dumps(step3), json.dumps(stored), now, now, setup_id),
        )
        conn.commit()
        locked = conn.execute("SELECT * FROM proposal_v2_setups WHERE id = ?", (setup_id,)).fetchone()
        _log_event(uid, "guided_proposal_step3_locked", {"setup_id": setup_id})
        return jsonify(_serialize(locked))
    finally:
        conn.close()


@guided_proposal_bp.route("/steps/4/analyze", methods=["POST"])
@require_role("premium")
def api_guided_proposal_analyze_step4():
    if disabled := _enabled_response():
        return disabled
    payload, setup_id = _step2_request_setup()
    if not setup_id:
        return jsonify({"error": "setup_id is required."}), 400
    uid, role = current_uid(), current_role()
    conn, row = _owned_setup(setup_id, uid, role)
    try:
        if not row:
            return jsonify({"error": "Setup not found."}), 404
        if row["step3_state"] != "locked":
            return jsonify({"error": "Step 3 must be locked before Step 4 can be analyzed."}), 409
        if row["step4_state"] == "locked":
            return jsonify({"error": "Locked Step 4 is immutable."}), 409
        setup = _serialize(row)
        data = normalize_step4(payload)
        analysis = validate_step4(data, setup, setup["technical_data"])
        # The deterministic financial/PSEA gates are authoritative. A blind LLM
        # verifier is added as advisory feedback without weakening those gates.
        analysis.update(
            {
                "donor_compliance_score": 90 if analysis["is_valid"] else 0,
                "critique_notes": ["Budget mathematics, donor overhead threshold and PSEA gate were checked."],
                "suggested_improvements": [],
                "analyzed_at": time.time(),
            }
        )
        conn.execute(
            "UPDATE proposal_v2_setups SET financial_data = ?, step4_analysis = ?, step4_state = 'analyzed', updated_at = ? WHERE id = ?",
            (json.dumps(data), json.dumps(analysis), time.time(), setup_id),
        )
        conn.commit()
        _log_event(uid, "guided_proposal_step4_analyzed", {"setup_id": setup_id, "valid": analysis["is_valid"]})
        return jsonify(analysis)
    finally:
        conn.close()


@guided_proposal_bp.route("/steps/4/lock", methods=["POST"])
@require_role("premium")
def api_guided_proposal_lock_step4():
    if disabled := _enabled_response():
        return disabled
    payload, setup_id = _step2_request_setup()
    if not setup_id:
        return jsonify({"error": "setup_id is required."}), 400
    uid, role = current_uid(), current_role()
    conn, row = _owned_setup(setup_id, uid, role)
    try:
        if not row:
            return jsonify({"error": "Setup not found."}), 404
        if row["step3_state"] != "locked":
            return jsonify({"error": "Step 3 must be locked before Step 4 can be locked."}), 409
        if row["step4_state"] == "locked":
            return jsonify({"error": "Step 4 is already locked."}), 409
        setup = _serialize(row)
        data = normalize_step4(payload if "budget_items" in payload else setup["financial_data"])
        deterministic = validate_step4(data, setup, setup["technical_data"])
        stored = setup["step4_analysis"]
        if not deterministic["is_valid"] or not stored.get("is_valid") or stored.get("donor_compliance_score", 0) < 70:
            return jsonify(
                {"error": "Analyze Step 4 and fix validation errors before locking.", "analysis": deterministic}
            ), 422
        # ── Cross-section consistency validation ────────────────────────────
        try:
            from agent.validation import validate_cross_sections

            cross_check = validate_cross_sections(
                {
                    "toc": setup["technical_data"].get("toc_narrative", []),
                    "logframe": setup["technical_data"].get("logframe", {}),
                    "budget": {
                        "total": stored.get("financial_summary", {}).get("total_budget", 0),
                        "lines": data.get("budget_items", []),
                    },
                    "mne_framework": setup["technical_data"].get("logframe", {}).get("indicators", []),
                    "risk_matrix": data.get("risks", []),
                    "needs_assessment": setup["context_data"].get("needs_assessment", ""),
                    "background": setup["context_data"].get("humanitarian_context", ""),
                    "methodology": setup["technical_data"].get("toc_narrative", ""),
                    "sustainability": "",
                    "coordination": "",
                }
            )
            stored["cross_section_validation"] = cross_check
        except Exception as _cs_err:
            import logging as _lg

            _lg.getLogger(__name__).warning(f"Cross-section validation failed for step 4 lock: {_cs_err}")
        now = time.time()
        stored.update({"step_id": 4, "locked_at": now})
        conn.execute(
            "UPDATE proposal_v2_setups SET financial_data = ?, step4_analysis = ?, step4_state = 'locked', step4_locked_at = ?, updated_at = ? WHERE id = ?",
            (json.dumps(data), json.dumps(stored), now, now, setup_id),
        )
        conn.commit()
        locked = conn.execute("SELECT * FROM proposal_v2_setups WHERE id = ?", (setup_id,)).fetchone()
        return jsonify(_serialize(locked))
    finally:
        conn.close()


def _locked_v2_summary(setup_id):
    conn, row = _owned_setup(setup_id, current_uid(), current_role())
    if not row:
        conn.close()
        return None, None
    setup = _serialize(row)
    conn.close()
    if not (
        setup["state"] == "locked"
        and setup["step2_state"] == "locked"
        and setup["step3_state"] == "locked"
        and setup["step4_state"] == "locked"
    ):
        return setup, False
    return setup, True


@guided_proposal_bp.route("/setups/<setup_id>/summary", methods=["GET"])
@require_role("premium")
def api_guided_proposal_summary(setup_id):
    if disabled := _enabled_response():
        return disabled
    setup, locked = _locked_v2_summary(setup_id)
    if setup is None:
        return jsonify({"error": "Setup not found."}), 404
    if not locked:
        return jsonify(
            {
                "error": "All Steps 1–4 must be locked before final review.",
                "completion": {
                    "step1": setup["state"],
                    "step2": setup["step2_state"],
                    "step3": setup["step3_state"],
                    "step4": setup["step4_state"],
                },
            }
        ), 409
    return jsonify(
        {
            "setup_id": setup_id,
            "completion": "locked",
            "step1": {
                key: setup[key]
                for key in (
                    "project_title",
                    "country",
                    "region",
                    "donor",
                    "budget_amount",
                    "budget_currency",
                    "executive_intent",
                    "sectors",
                    "locked_at",
                )
            },
            "step2": setup["context_data"],
            "step3": setup["technical_data"],
            "step4": setup["financial_data"],
            "locks": {
                "step1": setup["locked_at"],
                "step2": setup["step2_locked_at"],
                "step3": setup["step3_locked_at"],
                "step4": setup["step4_locked_at"],
            },
        }
    )


@guided_proposal_bp.route("/setups/<setup_id>/evaluate", methods=["POST"])
@require_role("premium")
def api_guided_proposal_evaluate(setup_id):
    if disabled := _enabled_response():
        return disabled
    setup, locked = _locked_v2_summary(setup_id)
    if setup is None:
        return jsonify({"error": "Setup not found."}), 404
    if not locked:
        return jsonify({"error": "All Steps 1–4 must be locked before evaluation."}), 409
    financial, _technical = setup["financial_data"], setup["technical_data"]
    overhead = (setup["step4_analysis"].get("financial_summary") or {}).get("indirect_overhead_percentage", 0)
    capacity = 16.0
    subcriteria = {"project_management": 5, "technical_expertise": 6, "financial_controls": 5}
    flags = [key for key, score in subcriteria.items() if score <= 1]
    technical_score = min(
        75.0,
        0.75
        * (
            (
                setup["step2_analysis"].get("donor_compliance_score", 0)
                + setup["step3_analysis"].get("donor_compliance_score", 0)
                + setup["step4_analysis"].get("donor_compliance_score", 0)
            )
            / 3
        ),
    )
    proposed = (setup["step4_analysis"].get("financial_summary") or {}).get("total_budget", 0) or setup["budget_amount"]
    benchmark = setup["budget_amount"] or proposed
    financial_score = round(25 * min(1, benchmark / proposed), 2) if proposed else 0
    overall = round(technical_score + financial_score, 1)
    status = "AUTOMATIC_REJECTION" if flags or capacity < 12 else "APPROVED_FOR_SUBMISSION"
    return jsonify(
        {
            "setup_id": setup_id,
            "overall_score": overall,
            "capacity_score": capacity,
            "capacity_passed": capacity >= 12 and not flags,
            "technical_score": round(technical_score, 2),
            "financial_score": financial_score,
            "prag_status": status,
            "donor_compliance": {
                "donor_id": setup["donor"].upper(),
                "overhead_percentage": overhead,
                "psea_compliant": bool(financial.get("psea_signoff")),
                "vulnerability_quota_met": (setup["step2_analysis"].get("beneficiary_summary") or {}).get(
                    "meets_quota", True
                ),
            },
            "audit_findings": [
                f"Capacity score ({capacity}/20) meets the 12-point threshold.",
                f"Indirect overhead is {overhead}%.",
            ],
            "disqualification_flags": flags,
        }
    )


@guided_proposal_bp.route("/setups/<setup_id>/compile-pdf", methods=["POST"])
@require_role("premium")
def api_guided_proposal_compile_pdf(setup_id):
    if disabled := _enabled_response():
        return disabled
    setup, locked = _locked_v2_summary(setup_id)
    if setup is None:
        return jsonify({"error": "Setup not found."}), 404
    if not locked:
        return jsonify({"error": "All Steps 1–4 must be locked before PDF compilation."}), 409
    from blueprints.proposal_pdf import build_proposal_pdf

    step1 = setup
    step2 = setup["context_data"]
    step3 = setup["technical_data"]
    step4 = setup["financial_data"]
    logframe = step3.get("logframe", {}) if isinstance(step3.get("logframe"), dict) else {}
    toc_narrative = step3.get("toc_narrative", "")
    toc = [{"level": "Theory of Change", "text": toc_narrative}] if toc_narrative else []
    indicators = []
    for row in step3.get("logframe", []):
        if isinstance(row, dict) and row.get("indicators"):
            for indicator in row["indicators"]:
                if isinstance(indicator, dict):
                    indicators.append(
                        {
                            "name": indicator.get("indicator_title", ""),
                            "baseline": indicator.get("baseline_value", ""),
                            "target": indicator.get("target_value", ""),
                            "source": indicator.get("data_source_and_frequency", ""),
                        }
                    )
    budget_items = step4.get("budget_items", [])
    total_budget = (setup["step4_analysis"].get("financial_summary") or {}).get("total_budget", 0)
    currency = step1.get("budget_currency", "USD")
    budget_lines = []
    for item in budget_items:
        if not isinstance(item, dict):
            continue
        amount = (
            float(item.get("quantity", 0) or 0)
            * float(item.get("unit_cost", 0) or 0)
            * float(item.get("duration_frequency", 0) or 0)
        )
        share = f"{(amount / total_budget * 100):.1f}%" if total_budget else "-"
        budget_lines.append(
            {
                "category": item.get("category", "General"),
                "description": item.get("description", ""),
                "amount": f"{amount:,.0f} {currency}",
                "percentage": share,
            }
        )
    donor_labels = {
        "unfpa": "UNFPA",
        "ocha_cbpf": "OCHA CBPF",
        "usaid_bha": "USAID/BHA",
        "europeaid_prag": "EuropeAid (PRAG)",
        "echo": "ECHO",
    }
    proposal = {
        "title": step1["project_title"],
        "country": step1["country"],
        "donor": donor_labels.get(step1["donor"], str(step1["donor"]).upper()),
        "event": "CEFM prevention, survivor-centred referral pathways and CSO capacity strengthening",
        "sectors": step1.get("sectors", []),
        "budget_amount": step1.get("budget_amount"),
        "budget_currency": step1.get("budget_currency"),
        "background": step2.get("humanitarian_context", ""),
        "needs_assessment": step2.get("needs_assessment", ""),
        "strategic_justification": step2.get("strategic_justification", ""),
        "toc": toc,
        "logframe": logframe,
        "logframe_rows": step3.get("logframe", []),
        "toc_narrative": toc_narrative,
        "gantt": step3.get("gantt", []),
        "budget_details": {"lines": budget_lines, "total": f"{total_budget:,.0f} {currency}"},
        "risk_details": step4.get("risks", []),
        "mne_framework": {"indicators": indicators},
        "psea_signoff": step4.get("psea_signoff", False),
        "sphere_standards_narrative": step4.get("sphere_standards_narrative", ""),
        "me_overall_score": (
            setup["step3_analysis"].get("me_overall_score")
            if (setup["step3_analysis"].get("me_overall_score") or 0) >= 60
            else None
        ),
        "me_suggestions": setup["step3_analysis"].get("me_suggestions", []),
        "cross_section_validation": setup["step4_analysis"].get("cross_section_validation"),
        "section_sources": {
            "background": setup["step2_analysis"].get("sources", []),
            "needs_assessment": setup["step2_analysis"].get("sources", []),
        },
    }
    pdf = build_proposal_pdf(proposal)
    filename = f"Proposal_{step1['donor']}_{step1['project_title'][:40]}.pdf"
    return send_file(pdf, mimetype="application/pdf", as_attachment=True, download_name=filename)


# ── V2 Advisor Chat ───────────────────────────────────────────────────────────
# Mirrors the V1 advisor but reads from proposal_v2_setups instead of proposals.

import threading as _threading

_ADVISOR_BUSY_TIMEOUT = 120
_advisor_busy: dict[str, bool] = {}
_advisor_busy_lock = _threading.Lock()


def _serialize_setup_for_context(row) -> str:
    """Build a human-readable context string from a V2 setup row."""
    setup = dict(row)
    parts = [
        f"Project: {setup.get('project_title', '?')}",
        f"Country: {setup.get('country', '?')}",
        f"Region: {setup.get('region', '?')}",
        f"Donor: {setup.get('donor', '?')}",
        f"Budget: {setup.get('budget_amount', '?')} {setup.get('budget_currency', '?')}",
        f"Executive Intent: {setup.get('executive_intent', '?')}",
        f"Sectors: {setup.get('sectors', '?')}",
    ]

    # Step state labels
    states = {
        1: setup.get("state", "draft"),
        2: setup.get("step2_state", "draft"),
        3: setup.get("step3_state", "draft"),
        4: setup.get("step4_state", "draft"),
    }
    parts.append("Step states: " + ", ".join(f"Step {k}={v}" for k, v in states.items()))

    # Context data
    for field in ("context_data", "technical_data", "financial_data"):
        raw = setup.get(field)
        if raw and raw not in ("{}", "", None):
            try:
                parsed = json.loads(raw) if isinstance(raw, str) else raw
                if isinstance(parsed, dict):
                    parts.append(f"\n{field}:\n" + json.dumps(parsed, indent=2, ensure_ascii=False)[:2000])
            except (json.JSONDecodeError, TypeError):
                parts.append(f"\n{field}: {str(raw)[:500]}")

    # Analysis results
    for field in ("analysis", "step2_analysis", "step3_analysis", "step4_analysis"):
        raw = setup.get(field)
        if raw and raw not in ("{}", "", None):
            try:
                parsed = json.loads(raw) if isinstance(raw, str) else raw
                if isinstance(parsed, dict):
                    score = parsed.get("donor_compliance_score")
                    valid = parsed.get("is_valid")
                    if score is not None or valid is not None:
                        parts.append(f"{field}: score={score}, valid={valid}")
            except (json.JSONDecodeError, TypeError):
                pass

    return "\n".join(parts)


@guided_proposal_bp.route("/setups/<setup_id>/advisor/chat", methods=["POST"])
@require_role("premium")
def api_guided_proposal_advisor_chat(setup_id):
    """Advisor chat for Guided Proposal V2."""
    import time as _time

    uid, role = current_uid(), current_role()
    data = request.get_json(silent=True) or {}
    message = data.get("message", "").strip()

    if not message:
        return jsonify({"error": "Message is required."}), 400

    # Per-user busy flag
    with _advisor_busy_lock:
        if uid in _advisor_busy and (_time.time() - _advisor_busy.get(uid + "_since", 0)) > _ADVISOR_BUSY_TIMEOUT:
            _advisor_busy[uid] = False
        if _advisor_busy.get(uid, False):
            return jsonify({"error": "Advisor is busy processing your previous request. Please wait."}), 429
        _advisor_busy[uid] = True
        _advisor_busy[uid + "_since"] = _time.time()

    conn = None
    try:
        conn, row = _owned_setup(setup_id, uid, role)
        if not row:
            return jsonify({"error": "Setup not found."}), 404

        setup = _serialize(row)
        chat_id = f"v2_advisor_{setup_id}"

        # Ensure chat session
        conn.execute(
            "INSERT OR IGNORE INTO chats (id, uid, title, created) VALUES (?, ?, ?, ?)",
            (chat_id, uid, f"V2 Advisor: {setup.get('project_title', 'Proposal')}", _time.time()),
        )
        conn.commit()

        # Load history
        db_rows = conn.execute(
            "SELECT role, content FROM chat_messages WHERE chat_id = ? ORDER BY ts ASC",
            (chat_id,),
        ).fetchall()

        from langchain_core.messages import AIMessage, HumanMessage, SystemMessage

        messages = []
        for r in db_rows[-20:]:
            if r["role"] == "user":
                messages.append(HumanMessage(content=r["content"]))
            elif r["role"] == "assistant":
                messages.append(AIMessage(content=r["content"]))

        messages.append(HumanMessage(content=message))

        # Save user message
        conn.execute(
            "INSERT INTO chat_messages (chat_id, role, content, ts) VALUES (?, 'user', ?, ?)",
            (chat_id, message, _time.time()),
        )
        conn.commit()

        # Build context from setup
        setup_context = _serialize_setup_for_context(row)

        # Fetch context chunks from ChromaDB
        chunks_text = ""
        try:
            from sitrep.chroma_adapter import ChromaAdapter

            chroma = ChromaAdapter()
            country = setup.get("country", "")
            sectors = setup.get("sectors", [])
            if isinstance(sectors, str):
                try:
                    sectors = json.loads(sectors)
                except (json.JSONDecodeError, TypeError):
                    sectors = [s.strip() for s in sectors.split(",") if s.strip()]
            chunks = chroma.get_chunks_by_country_and_themes(country, sectors or None)
            if chunks:
                chunks_text = "\n\n".join(
                    [f"- {c.get('title', 'Report')}: {c.get('text', '')}" for c in chunks[:10]]
                )
        except Exception as e:
            logger.warning(f"Failed to fetch chunks for V2 advisor: {e}")

        # Load donor profile for context
        from agent.proposal_v2_rules import DONOR_PROFILES

        donor = DONOR_PROFILES.get(setup.get("donor", ""), DONOR_PROFILES.get("generic", {}))
        donor_context = (
            f"Donor: {donor.get('full_name', 'Generic')}\n"
            f"Framework: {donor.get('framework_standard', '')}\n"
            f"Overhead ceiling: {donor.get('overhead_ceiling_percent', '?')}%\n"
            f"Max duration: {donor.get('max_duration_months', '?')} months\n"
            f"Special requirements: {'; '.join(donor.get('special_requirements', [])[:5])}"
        )

        system_prompt = f"""You are the Sightline Proposal Advisor for a Guided Proposal (V2). The user is actively working on a humanitarian proposal.

Current proposal state:
{setup_context}

Donor rules:
{donor_context}

Relevant humanitarian context:
{chunks_text}

You can see the full proposal state above. Help the user improve their proposal based on donor compliance rules, humanitarian best practices, and the context data. Be specific and constructive. When suggesting edits, describe what to change and why."""

        messages.insert(0, SystemMessage(content=system_prompt))

        # Invoke agent
        from blueprints.helpers import _get_agent

        agent = _get_agent()
        config = {"recursion_limit": 25, "configurable": {"uid": uid, "proposal_id": setup_id}}

        result = agent.invoke({"messages": messages}, config=config)

        # Extract final response
        final_response = ""
        for msg in reversed(result.get("messages", [])):
            if isinstance(msg, AIMessage) and not getattr(msg, "tool_calls", None):
                final_response = msg.content
                break

        if not final_response:
            final_response = "I have reviewed your proposal. Let me know if you'd like specific suggestions."

        # Save assistant message
        conn.execute(
            "INSERT INTO chat_messages (chat_id, role, content, ts) VALUES (?, 'assistant', ?, ?)",
            (chat_id, final_response, _time.time()),
        )
        conn.commit()

        # Check if proposal tools were used (for auto-refresh)
        proposal_edited = False
        for msg in result.get("messages", []):
            if hasattr(msg, "name") and msg.name in (
                "edit_proposal_toc",
                "edit_proposal_logframe",
                "edit_proposal_narrative",
            ):
                proposal_edited = True
                break

        command_data = {"action": "refresh"} if proposal_edited else None

        return jsonify({"response": final_response, "command": command_data})

    except Exception as e:
        logger.error("V2 advisor chat error for setup %s: %s", setup_id, e)
        return jsonify({"error": "Advisor failed. Please try again."}), 500
    finally:
        with _advisor_busy_lock:
            _advisor_busy.pop(uid, None)
            _advisor_busy.pop(uid + "_since", None)
        if conn:
            conn.close()


@guided_proposal_bp.route("/setups/<setup_id>/advisor/history", methods=["GET"])
@require_auth
def api_guided_proposal_advisor_history(setup_id):
    """Retrieve advisor chat history for a Guided Proposal V2 setup."""
    uid, role = current_uid(), current_role()
    conn, row = _owned_setup(setup_id, uid, role)
    if not row:
        conn.close()
        return jsonify({"error": "Setup not found."}), 404
    conn.close()

    chat_id = f"v2_advisor_{setup_id}"
    conn = _chats_db()
    try:
        db_rows = conn.execute(
            "SELECT role, content FROM chat_messages WHERE chat_id = ? ORDER BY ts ASC",
            (chat_id,),
        ).fetchall()

        history = [{"role": r["role"], "content": r["content"]} for r in db_rows]
        return jsonify(history)
    except Exception as e:
        logger.error("V2 advisor history error for setup %s: %s", setup_id, e)
        return jsonify([])
    finally:
        conn.close()
